# -*- coding: utf-8 -*-
"""
Enhanced competition paper for ntFAST project
~35-40 pages, Kazakh language, 3rd person
With screenshots and detailed anti-fraud section
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ============================================================
# PAGE SETUP: A4, margins 30-15-20-20
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Footer with page numbers
footer = doc.sections[0].footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
run._element.append(fldChar1)
run2 = fp.add_run()
instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w')))
run2._element.append(instrText)
run3 = fp.add_run()
fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
run3._element.append(fldChar2)

style = doc.styles['Normal']
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 1.5
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

# Fix font for East Asian
rPr = style.element.get_or_add_rPr()
rFonts_elem = parse_xml(
    r'<w:rFonts {} w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
    r'w:cs="Times New Roman" w:eastAsia="Times New Roman"/>'.format(nsdecls('w'))
)
rPr.append(rFonts_elem)

SCREENSHOTS_DIR = r"C:\Users\Admin\Desktop\screenshots"

def add_empty_lines(n=1):
    for _ in range(n):
        p = doc.add_paragraph('')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

def add_centered(text, bold=False, size=14, caps=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text.upper() if caps else text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_normal(text, bold=False, indent_first=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p

def add_bold_and_normal(bold_text, normal_text, indent_first=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r1 = p.add_run(bold_text)
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r2 = p.add_run(normal_text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    return p

def add_heading_custom(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_list_item(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run('\u2013 ' + text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p

def add_image(filename, caption, width_cm=15):
    img_path = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        r = cap.add_run(caption)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.italic = True
    else:
        add_normal('[Сурет: ' + caption + ' - файл табылмады: ' + filename + ']')

def add_formula(text, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.italic = True
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

def set_cell(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)

def shade_cells(row, color="D5E8F0"):
    for cell in row.cells:
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shading)

# ============================================================
# TITLE PAGE
# ============================================================
add_empty_lines(3)
add_centered('СТУДЕНТТЕРДIҢ ҒЫЛЫМИ-ЗЕРТТЕУ ЖҰМЫСТАРЫНЫҢ', bold=True, size=14)
add_centered('РЕСПУБЛИКАЛЫҚ КОНКУРСЫ', bold=True, size=14)
add_empty_lines(3)
add_centered('Девиз: \u00ABDIGITAL GUARDIAN\u00BB', bold=True, size=16)
add_empty_lines(2)
add_centered('Секция: Ақпараттық технологиялар', bold=False, size=14)
add_empty_lines(1)
add_centered('Тақырыбы:', bold=True, size=14)
add_centered('\u00ABҚАРЖЫЛЫҚ ТРАНЗАКЦИЯЛАРДЫ ИНТЕЛЛЕКТУАЛДЫ', bold=True, size=16)
add_centered('ТАЛДАУ ЖҮЙЕСIН ӘЗIРЛЕУ\u00BB', bold=True, size=16)
add_centered('(ntFAST \u2014 Financial Analysis System for Transactions)', bold=False, size=13)
add_empty_lines(10)
doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_custom('МАЗМҰНЫ')
add_empty_lines(1)

toc_items = [
    ('Кiрiспе', '3'),
    ('1 Қаржылық транзакцияларды талдаудың теориялық негiздерi', '6'),
    ('  1.1 Қаржылық алаяқтық және оны анықтау мәселесi', '6'),
    ('  1.2 Банк үзiндiлерiн талдаудың қазiргi жағдайы', '8'),
    ('  1.3 Жасанды интеллект пен машиналық оқыту әдiстерiн қолдану', '10'),
    ('  1.4 Қазақстандағы қаржылық технологиялар нарығы', '12'),
    ('2 ntFAST жүйесiнiң архитектурасы мен технологиялық шешiмдерi', '14'),
    ('  2.1 Жүйенiң жалпы архитектурасы', '14'),
    ('  2.2 Backend: FastAPI және Python технологиялары', '16'),
    ('  2.3 Frontend: React және TypeScript', '18'),
    ('  2.4 Деректер базасы: PostgreSQL', '19'),
    ('  2.5 Антифрод-талдау модульдерi (FraudEngine v4)', '20'),
    ('  2.6 Банк үзiндiлерiн тану жүйесi', '25'),
    ('  2.7 Нақты уақыт режимiндегi WebSocket байланыс', '26'),
    ('3 Жүйенi әзiрлеу және тестiлеу нәтижелерi', '27'),
    ('  3.1 Жүйенiң функционалдық мүмкiндiктерi', '27'),
    ('  3.2 Антифрод-талдау нәтижелерi', '30'),
    ('  3.3 Қауiпсiздiк және рөлге негiзделген қол жеткiзу', '32'),
    ('  3.4 Көп тiлдi интерфейс және адаптивтi дизайн', '33'),
    ('  3.5 Тестiлеу нәтижелерi', '34'),
    ('Қорытынды', '35'),
    ('Пайдаланылған әдебиеттер тiзiмi', '37'),
]

for title, page in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    dots_count = max(3, 65 - len(title))
    run = p.add_run(title + ' ' + '.' * dots_count + ' ' + page)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

doc.add_page_break()

# ============================================================
# INTRODUCTION (pages 3-5)
# ============================================================
add_heading_custom('КIРIСПЕ')
add_empty_lines(1)

add_bold_and_normal('Зерттеу жұмысының өзектiлiгi. ',
    'Қазiргi заманда цифрлық қаржылық жүйелердiң қарқынды дамуы қаржылық транзакциялар көлемiнiң '
    'айтарлықтай артуына алып келдi. Қазақстан Республикасында электрондық төлемдер нарығы жыл сайын '
    '30\u201340%-ға өсуде, бұл ретте қаржылық алаяқтық пен заңсыз ақша айналымы жағдайлары да '
    'көбеюде. Қазақстан Республикасы Ұлттық Банкiнiң мәлiметтерi бойынша, 2024 жылы электрондық '
    'төлем операцияларының жалпы көлемi 50 триллион теңгеден асты, ал 2025 жылы бұл көрсеткiш '
    '65 триллион теңгеге жеттi. Kaspi Bank экожүйесiнде күнделiктi 15 миллионнан астам транзакция '
    'жасалады, Halyk Bank жүйесiнде \u2014 8 миллионнан астам.')

add_normal('Қаржылық алаяқтық мәселесi жаһандық деңгейде өзектi. Association of Certified '
    'Fraud Examiners (ACFE) ұйымының 2024 жылғы есебiне сәйкес, ұйымдар жылдық табысының '
    'шамамен 5%-ын алаяқтыққа жоғалтады, бұл жаһандық деңгейде 4,7 триллион АҚШ долларын құрайды. '
    'Қазақстан Республикасы Қаржылық мониторинг агенттiгiнiң мәлiметтерi бойынша, 2024 жылы '
    '50 мыңнан астам күдiктi қаржылық операция тiркелдi, ал электрондық алаяқтық жағдайлары '
    'жыл сайын 25\u201330%-ға артуда.')

add_normal('Дәстүрлi тәсiлде банк үзiндiлерiн талдау қолмен жүргiзiледi, бұл процесс көп '
    'уақытты қажет етедi және адам факторына байланысты қателiктерге жол берiледi. Бiр банк '
    'үзiндiсiн қолмен талдау 2\u20138 сағатты алады, ал күрделi жағдайларда \u2014 бiрнеше '
    'жұмыс күнiн. Сонымен қатар, қолмен талдау кезiнде күрделi алаяқтық схемаларды (транзакцияларды '
    'бөлшектеу, айналмалы аударымдар, түнгi операциялар сериясы) анықтау қиынға соғады. '
    'Осы мәселелердi шешу мақсатында ntFAST (Financial Analysis System for Transactions) \u2014 '
    'қаржылық транзакцияларды интеллектуалды талдау жүйесi әзiрлендi.')

add_bold_and_normal('Зерттеу жұмысының мақсаты \u2014 ',
    'Қазақстан банктерiнiң үзiндiлерiн автоматты түрде талдайтын, 11 модульдi антифрод-жүйесi бар, '
    'нақты уақыт режимiнде жұмыс iстейтiн, жасанды интеллект элементтерiн қолданатын интеллектуалды '
    'қаржылық талдау жүйесiн әзiрлеу.')

add_normal('Зерттеу жұмысының мiндеттерi:', bold=True)

tasks = [
    'Қаржылық транзакцияларды талдау саласындағы бар шешiмдердi зерттеу және салыстырмалы талдау жүргiзу;',
    'Қазақстан банктерiнiң (Kaspi, Halyk, Jusan, Forte, BCC) PDF үзiндiлерiн автоматты тану алгоритмiн әзiрлеу;',
    '11 модульден тұратын FraudEngine v4 антифрод-талдау жүйесiн жобалау және iске асыру;',
    'FastAPI + React технологияларына негiзделген клиент-сервер архитектурасын әзiрлеу;',
    'WebSocket арқылы нақты уақыт режимiндегi талдау прогресiн бақылау жүйесiн құру;',
    'Рөлге негiзделген қол жеткiзу (RBAC) жүйесiн және қауiпсiздiк модулiн iске асыру;',
    'Жүйенi тестiлеу, антифрод-модульдердiң дәлдiгiн бағалау және нәтижелердi верификациялау.',
]
for task in tasks:
    add_list_item(task)

add_bold_and_normal('Зерттеу объектiсi \u2014 ', 'банк үзiндiлерiн автоматтандырылған талдау және қаржылық алаяқтықты анықтау процесi.')
add_bold_and_normal('Зерттеу пәнi \u2014 ', 'қаржылық транзакцияларды интеллектуалды талдау жүйесiн әзiрлеу технологиялары мен әдiстерi.')

add_normal('Зерттеудiң ғылыми жаңалығы мынада:', bold=True)
novelty = [
    'Қазақстан банктерiнiң (Kaspi, Halyk, Jusan, Forte, BCC) PDF үзiндiлерiн автоматты тану мен талдау жүйесi алғаш рет әзiрлендi \u2014 BankDetector модулi үш деңгейлi анықтау алгоритмiн (keywords matching, IBAN pattern analysis, structural hints) қолданады;',
    'Аккаунт профилiне негiзделген контекстуалды салмақтау (contextual weighting) арқылы composite risk score есептейтiн 11 модульдi FraudEngine v4 антифрод-жүйесi ұсынылды \u2014 AccountProfiler модулi аккаунт типiн автоматты анықтап, PatternWhitelist арқылы false positive көрсеткiшiн 15%-ға дейiн төмендетедi;',
    'Rule-based тәсiлмен нақты уақыт режимiнде жұмыс iстейтiн, WebSocket прогресс-бақылау жүйесiмен интеграцияланған, графтық талдау (NetworkX) элементтерiн қамтитын толық стектi веб-қосымша әзiрлендi.',
]
for n in novelty:
    add_list_item(n)

add_bold_and_normal('Практикалық маңыздылығы. ',
    'Әзiрленген ntFAST жүйесi қаржылық мекемелер, аудиторлық компаниялар, қаржылық бақылау органдары '
    'және жеке тұлғалар үшiн банк үзiндiлерiн талдау процесiн айтарлықтай жеңiлдетедi. Жүйе '
    'бiрнеше сағатты алатын қолмен талдау процесiн 2\u20135 минутқа қысқартады. 11 антифрод-модуль '
    'арқылы күрделi алаяқтық схемаларын автоматты түрде анықтайды. Жүйе қазақ, орыс және ағылшын '
    'тiлдерiнде жұмыс iстейдi, адаптивтi дизайнмен қамтамасыз етiлген.')

add_bold_and_normal('Зерттеу әдiстемесi ',
    'жүйелiк тәсiлге негiзделдi. Алдымен бар шешiмдерге салыстырмалы талдау жүргiзiлдi, '
    'содан кейiн жүйенiң талаптар спецификациясы құрастырылды, объектiге бағытталған жобалау '
    'жүргiзiлдi, жүйе итерациялық түрде әзiрлендi және тестiлендi. Agile әдiснамасы (Scrum) '
    'қолданылды \u2014 екi апталық спринттер, тұрақты ревьюлар мен ретроспективалар. '
    'Зерттеу барысында индукция, дедукция, моделдеу, эксперимент және салыстырмалы талдау '
    'әдiстерi қолданылды.')

add_normal('Зерттеу жұмысы кiрiспеден, үш тараудан, қорытындыдан және пайдаланылған әдебиеттер '
    'тiзiмiнен тұрады. Бiрiншi тарауда қаржылық транзакцияларды талдаудың теориялық негiздерi, '
    'қаржылық алаяқтықты анықтау әдiстерi мен Қазақстандағы FinTech нарығының қазiргi жағдайы '
    'қарастырылды. Екiншi тарауда ntFAST жүйесiнiң архитектурасы, қолданылған технологиялар мен '
    '11 антифрод-модульдiң жұмыс принциптерi толық сипатталды. Үшiншi тарауда жүйенiң функционалдық '
    'мүмкiндiктерi, тестiлеу нәтижелерi мен практикалық қолдану мәселелерi талданды.')

doc.add_page_break()

# ============================================================
# CHAPTER 1: THEORETICAL FOUNDATIONS (pages 6-13)
# ============================================================
add_heading_custom('1 ҚАРЖЫЛЫҚ ТРАНЗАКЦИЯЛАРДЫ ТАЛДАУДЫҢ ТЕОРИЯЛЫҚ НЕГIЗДЕРI')
add_empty_lines(1)

# 1.1
add_subheading('1.1 Қаржылық алаяқтық және оны анықтау мәселесi')

add_normal('Қаржылық алаяқтық \u2014 заңсыз жолмен қаржылық пайда табу мақсатында жасалатын алдау '
    'әрекеттерiнiң жиынтығы. Халықаралық тәжiрибеде қаржылық алаяқтықтың бiрнеше негiзгi түрлерi '
    'ажыратылады: кредиттiк карта алаяқтығы, жеке басты ұрлау (identity theft), ақша жылыстату '
    '(money laundering), транзакцияларды бөлшектеу (structuring немесе smurfing), сондай-ақ '
    'цифрлық төлем жүйелерiндегi алаяқтық.')

add_normal('Association of Certified Fraud Examiners (ACFE) ұйымының 2024 жылғы Occupational '
    'Fraud Report есебiне сәйкес, әлем бойынша ұйымдар жылдық табысының шамамен 5%-ын '
    'алаяқтыққа жоғалтады. Бұл жаһандық масштабта 4,7 триллион АҚШ долларын құрайды. Қаржылық '
    'алаяқтықты анықтау (Fraud Detection) саласы соңғы онжылдықта айтарлықтай дамыды және бүгiнгi '
    'таңда технологиялық шешiмдерсiз тиiмдi жұмыс iстеу мүмкiн емес.')

add_normal('Қаржылық алаяқтықты анықтаудың негiзгi тәсiлдерi зерттеу барысында жiктелдi:')

approaches = [
    'Rule-based жүйелер \u2014 алдын ала анықталған ережелер негiзiнде күдiктi транзакцияларды табу. Бұл тәсiлдiң артықшылығы \u2014 түсiнiктi, бақылануы оңай, детерминистiк нәтижелер бередi. Кемшiлiгi \u2014 жаңа алаяқтық схемаларына автоматты бейiмделу қиын;',
    'Статистикалық әдiстер \u2014 транзакциялардың статистикалық сипаттамаларын (орташа мәнi, стандартты ауытқуы, медианасы) талдау арқылы аномалияларды анықтау;',
    'Машиналық оқыту (ML) әдiстерi \u2014 supervised (бақылаулы) және unsupervised (бақылаусыз) оқыту алгоритмдерi арқылы алаяқтық үлгiлерiн автоматты тану. Random Forest, Gradient Boosting, Neural Networks алгоритмдерi кеңiнен қолданылады;',
    'Гибридтi жүйелер \u2014 rule-based және ML тәсiлдерiн бiрiктiру арқылы жүйенiң тиiмдiлiгiн арттыру. ntFAST жүйесi осы категорияға жақын, себебi rule-based ядро графтық талдау элементтерiмен толықтырылған.',
]
for a in approaches:
    add_list_item(a)

add_normal('ntFAST жүйесiнде rule-based тәсiл таңдалды, себебi бұл тәсiл транзакциялардың нақты '
    'сипаттамаларына негiзделедi, нәтижелерi түсiнiктi және интерпретацияланатын (explainability), '
    'оқыту деректерiн талап етпейдi. Rule-based тәсiл қаржылық реттеушi органдардың талаптарына '
    '(compliance requirements) толық сәйкес келедi, себебi шешiм қабылдау процесi мөлдiр және '
    'аудитке жарамды. Бұл ерекшелiк Қазақстан Республикасының Қылмыстық жолмен алынған '
    'кiрiстердi жылыстатуға және терроризмдi қаржыландыруға қарсы iс-қимыл туралы Заңының '
    '(2009 ж.) талаптарына сәйкес маңызды.')

add_normal('Rule-based жүйелердiң негiзгi артықшылықтары егжей-тегжейлi талданды. Бiрiншiден, '
    'олар детерминистiк нәтижелер бередi \u2014 бiрдей транзакция әрқашан бiрдей нәтиже алады, '
    'бұл қаржылық аудит кезiнде маңызды. Екiншiден, domain expertise негiзiнде жаңа ережелердi '
    'жылдам қосу мүмкiндiгiн бередi \u2014 ML модельдерiн қайта оқытуды қажет етпейдi. Үшiншiден, '
    'жүйенiң өнiмдiлiгi ML inference-тан тезiрек \u2014 ережелердi тексеру O(n) күрделiлiкпен '
    'орындалады, мұндағы n \u2014 транзакциялар саны.')

add_normal('Алаяқтықтың негiзгi категориялары зерттеу барысында жiктелдi. Бiрiншi категория \u2014 '
    'транзакцияларды бөлшектеу (structuring/smurfing): iрi сомаларды мiндеттi мониторинг шегiнен '
    'төмен бiрнеше кiшiгiрiм аударымдарға бөлу. Қазақстан заңнамасы бойынша 1 000 000 теңгеден '
    'асатын қолма-қол ақша операциялары мiндеттi мониторингке жатады. Екiншi категория \u2014 '
    'айналмалы аударымдар (round-tripping): ақшаны бiрнеше аралық шот арқылы бастапқы шотқа '
    'қайтару. Үшiншi категория \u2014 жоғары тәуекелдi мерчанттермен операциялар (ойын '
    'платформалары, букмекерлiк кеңселер, криптовалюта биржалары). Төртiншi категория \u2014 '
    'уақыттық аномалиялар (түнгi iрi аударымдар, жиiлiктiң кенеттен артуы).')

add_normal('Қазақстан заңнамасына сәйкес, қаржылық мониторинг субъектiлерi күдiктi операцияларды '
    'анықтау және тиiстi органдарға хабарлау мiндетi бар. ntFAST жүйесi осы заңнамалық '
    'талаптарды ескере отырып жобаланды.')

# 1.2
add_subheading('1.2 Банк үзiндiлерiн талдаудың қазiргi жағдайы')

add_normal('Банк үзiндiсi (bank statement) \u2014 белгiлi бiр кезеңдегi шот бойынша жасалған '
    'барлық операциялардың толық есебi. Банк үзiндiлерi қаржылық аудит, салықтық тексеру, '
    'кредиттiк бағалау, тергеу iс-шаралары және due diligence процестерi барысында кеңiнен '
    'қолданылады.')

add_normal('Қазiргi таңда Қазақстанда банк үзiндiлерiн талдау негiзiнен қолмен жүргiзiледi. '
    'Аудиторлар мен қаржылық талдаушылар PDF немесе Excel форматындағы үзiндiлердi қолмен '
    'қарап шығады. Бұл процестiң негiзгi кемшiлiктерi:')

drawbacks = [
    'Уақыт шығыны \u2014 бiр үзiндiнi толық талдау 2\u20138 сағатты алады, күрделi жағдайларда бiрнеше жұмыс күнiн;',
    'Адам факторы \u2014 шаршаған немесе тәжiрибесiз маман маңызды детальдарды байқамай қалуы мүмкiн;',
    'Масштабтау мүмкiн еместiгi \u2014 бiрнеше мың транзакцияны қолмен талдау практикалық тұрғыдан мүмкiн емес;',
    'Стандарттау мәселесi \u2014 әр маман өзiнше талдайды, нәтижелердiң сәйкестiгi төмен;',
    'Күрделi схемаларды анықтау \u2014 транзакцияларды бөлшектеу, айналмалы аударымдар сияқты схемаларды қолмен анықтау қиын.',
]
for d in drawbacks:
    add_list_item(d)

add_normal('Нарықтағы бар шешiмдерге салыстырмалы талдау жүргiзiлдi. Нәтижелер 1-кестеде көрсетiлген.')
add_empty_lines(1)

add_normal('1-кесте \u2014 Бар шешiмдердiң салыстырмалы талдауы', bold=True, indent_first=False)
table1 = doc.add_table(rows=7, cols=6)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

h1 = ['Жүйе', 'ҚР банктерi', 'Қазақ тiлi', 'Антифрод модульдер', 'Нақты уақыт', 'Бағасы']
for i, h in enumerate(h1):
    set_cell(table1.rows[0].cells[i], h, bold=True, size=10)
shade_cells(table1.rows[0], "D5E8F0")

comp_data = [
    ['NICE Actimize', 'Жоқ', 'Жоқ', '20+', 'Иә', '$50K+/жыл'],
    ['Chainalysis', 'Iшiнара', 'Жоқ', '10+', 'Иә', '$30K+/жыл'],
    ['SAS Fraud Mgmt', 'Жоқ', 'Жоқ', '15+', 'Иә', '$100K+/жыл'],
    ['Elliptic', 'Жоқ', 'Жоқ', '8+', 'Иә', '$20K+/жыл'],
    ['Қолмен талдау', 'Иә', 'Иә', '0', 'Жоқ', 'Маман жалақысы'],
    ['ntFAST', 'Иә (5 банк)', 'Иә', '11', 'Иә (WS)', 'Тегiн/ашық код'],
]
for r_idx, row_data in enumerate(comp_data):
    for c_idx, val in enumerate(row_data):
        set_cell(table1.rows[r_idx+1].cells[c_idx], val, size=10)

add_empty_lines(1)

add_normal('Салыстырмалы талдау нәтижесiнде анықталды: халықаралық шешiмдер қымбат ($20K\u2013$100K+ '
    'жылына), Қазақстан банктерiн қолдамайды және жергiлiктi нарыққа бейiмделмеген. ntFAST '
    'жүйесi осы олқылықтарды толтыру мақсатында әзiрлендi.')

add_normal('Қазақстандағы негiзгi банктердiң үзiндi форматтары бiр-бiрiнен айтарлықтай ерекшеленедi. '
    'Kaspi Bank PDF үзiндiлерi графикалық элементтермен және категорияланған кестелiк деректермен '
    'ерекшеленсе, Halyk Bank үзiндiлерi мәтiндiк форматқа жақын. Jusan, Forte банктерiнiң де '
    'өзiндiк форматтары бар. Осы әртүрлi форматтарды бiрiңғай жүйемен талдау мәселесi ntFAST '
    'жүйесiн әзiрлеудiң негiзгi техникалық қиындықтарының бiрi болды.')

# 1.3
add_subheading('1.3 Жасанды интеллект пен машиналық оқыту әдiстерiн қолдану')

add_normal('Қаржылық транзакцияларды талдау саласында жасанды интеллект (AI) пен машиналық оқыту '
    '(ML) әдiстерi соңғы жылдары кеңiнен қолданылуда. Бұл әдiстердiң негiзгi артықшылығы \u2014 '
    'үлкен көлемдi деректердi жылдам өңдеу және адам көзiне байқалмайтын үлгiлердi (patterns) '
    'анықтау мүмкiндiгi.')

add_normal('Алаяқтықты анықтауда қолданылатын негiзгi AI/ML әдiстерi:')

ml_methods = [
    'Аномалия анықтау (Anomaly Detection) \u2014 Isolation Forest, Local Outlier Factor, Autoencoder алгоритмдерi транзакциялардың қалыпты профильден ауытқуын анықтайды;',
    'Классификация \u2014 Random Forest, Gradient Boosting (XGBoost, LightGBM), Neural Networks алгоритмдерi транзакцияларды заңды және алаяқтық деп жiктейдi;',
    'Графтық талдау (Graph Analysis) \u2014 транзакциялар желiсiн граф G=(V,E) ретiнде қарастырып, циклдердi, хабтарды және күдiктi кластерлердi анықтайды;',
    'Уақыттық қатарлар талдау (Time Series Analysis) \u2014 транзакциялардың уақыттық үлгiлерiн талдау, жиiлiктiң кенеттен артуын анықтау;',
    'NLP (Natural Language Processing) \u2014 транзакция сипаттамаларын мәтiндiк талдау арқылы категориялау және мерчант тәуекелiн бағалау.',
]
for m in ml_methods:
    add_list_item(m)

add_normal('ntFAST жүйесiнде rule-based тәсiл ML элементтерiмен толықтырылды. Атап айтқанда, '
    'TransactionGraphAnalyzer модулi транзакциялар желiсiн NetworkX кiтапханасы арқылы '
    'бағытталған граф ретiнде құрастырады. Графтық талдау әдiсi транзакциялар арасындағы '
    'құрылымдық байланыстарды анықтауда ерекше тиiмдi \u2014 циклдер (A\u2192B\u2192C\u2192A) '
    'ақша жылыстатудың белгiсi, хаб-түйiндер ақша жинақтау/тарату орталықтарын, тығыз '
    'кластерлер ұйымдасқан алаяқтық тобын көрсетуi мүмкiн.')

add_normal('Velocity Analyzer модулi уақыттық терезелер (sliding windows) тәсiлiн қолданады: '
    '1 сағаттық, 24 сағаттық және 7 күндiк терезелер арқылы транзакциялар жиiлiгi бақыланады.')

add_normal('Rule-based тәсiлдiң ML алдындағы артықшылығы \u2014 модельдi оқыту үшiн таңбаланған '
    '(labeled) деректер жиынтығы қажет емес, ал Қазақстан банктерiнiң транзакциялары бойынша '
    'мұндай деректер жиынтығы қолжетiмдi емес. Сонымен қатар, rule-based жүйе шешiмдердi '
    'түсiндiре алады (explainability), бұл қаржылық реттеу саласында маңызды талап.')

# 1.4
add_subheading('1.4 Қазақстандағы қаржылық технологиялар нарығы')

add_normal('Қазақстан Орталық Азиядағы қаржылық технологиялар (FinTech) нарығының көшбасшысы '
    'болып табылады. Astana International Financial Centre (AIFC) аясында құрылған FinTech Hub '
    '200-ден астам компанияны бiрiктiредi. Қазақстанда мобильдi банкинг пайдаланушыларының саны '
    '15 миллионнан асты, бұл елдiң ересек тұрғындарының 80%-дан астамын құрайды.')

add_normal('Kaspi Bank \u2014 Қазақстандағы ең iрi цифрлық экожүйе, оның Kaspi.kz платформасы '
    '13 миллионнан астам белсендi пайдаланушыға ие. Halyk Bank \u2014 елдегi ең iрi банк, оның '
    'Homebank қосымшасы 8 миллионнан астам пайдаланушыны қамтиды. Jusan Bank, Forte Bank және '
    'BCC банктерi де цифрлық трансформация процесiн белсендi жүргiзуде.')

add_normal('Қаржылық алаяқтық мәселесi Қазақстанда да өзектi. Iшкi iстер министрлiгiнiң '
    'мәлiметтерi бойынша, 2024 жылы интернет-алаяқтық бойынша 45 мыңнан астам өтiнiш тiркелдi. '
    'Электрондық алаяқтықтан келтiрiлген залал жылына 50 миллиард теңгеден асады.')

add_normal('ntFAST жүйесi осы нарықтық қажеттiлiкке жауап ретiнде әзiрлендi. Жүйе Қазақстан '
    'банктерiнiң нақты форматтарын қолдайды, қазақ және орыс тiлдерiнде жұмыс iстейдi, '
    'сондай-ақ жергiлiктi алаяқтық схемаларын ескередi. Жүйе ашық бастапқы код ретiнде '
    'қолжетiмдi, бұл оны коммерциялық шешiмдерге балама етедi.')

doc.add_page_break()

# ============================================================
# CHAPTER 2: ARCHITECTURE (pages 14-25)
# ============================================================
add_heading_custom('2 ntFAST ЖҮЙЕСIНIҢ АРХИТЕКТУРАСЫ МЕН ТЕХНОЛОГИЯЛЫҚ ШЕШIМДЕРI')
add_empty_lines(1)

# 2.1
add_subheading('2.1 Жүйенiң жалпы архитектурасы')

add_normal('ntFAST жүйесi клиент-сервер архитектурасына негiзделген заманауи веб-қосымша ретiнде '
    'әзiрлендi. Жүйе үш деңгейлi (three-tier) архитектураны қолданады: презентациялық деңгей '
    '(frontend), iскери логика деңгейi (backend) және деректер деңгейi (database).')

add_normal('Жүйенiң архитектуралық компоненттерi:')

components = [
    'Frontend (клиенттiк бөлiк) \u2014 React 18 + TypeScript негiзiнде SPA архитектурасымен әзiрлендi. Vite құралымен жинақталады;',
    'Backend (серверлiк бөлiк) \u2014 Python 3.11 + FastAPI фреймворкi, RESTful API архитектурасы, асинхронды ASGI өңдеу;',
    'Database (деректер базасы) \u2014 PostgreSQL 15, SQLAlchemy ORM, Alembic миграциялары;',
    'Task Queue (тапсырмалар кезегi) \u2014 Celery + Redis асинхронды тапсырмаларды өңдеу;',
    'WebSocket Server \u2014 нақты уақыт режимiндегi екi жақты байланыс, прогресс-бақылау.',
]
for c in components:
    add_list_item(c)

add_empty_lines(1)
add_normal('2-кесте \u2014 Технологиялық стек', bold=True, indent_first=False)
table2 = doc.add_table(rows=13, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

h2 = ['Компонент', 'Технология', 'Нұсқасы']
for i, h in enumerate(h2):
    set_cell(table2.rows[0].cells[i], h, bold=True, size=11)
shade_cells(table2.rows[0], "D5E8F0")

tech_data = [
    ['Backend фреймворк', 'FastAPI', '0.104+'],
    ['Бағдарламалау тiлi (backend)', 'Python', '3.11'],
    ['Frontend фреймворк', 'React', '18.2'],
    ['Бағдарламалау тiлi (frontend)', 'TypeScript', '5.0'],
    ['Build құралы', 'Vite', '5.0'],
    ['Деректер базасы', 'PostgreSQL', '15'],
    ['ORM', 'SQLAlchemy', '2.0'],
    ['Тапсырмалар кезегi', 'Celery + Redis', '5.3 / 7.0'],
    ['PDF парсинг', 'pdfplumber', '0.10'],
    ['Графтық талдау', 'NetworkX', '3.2'],
    ['Аутентификация', 'JWT (python-jose)', '3.3'],
    ['Көп тiлдi', 'i18next', '23.0'],
]
for r_idx, row_data in enumerate(tech_data):
    for c_idx, val in enumerate(row_data):
        set_cell(table2.rows[r_idx+1].cells[c_idx], val, size=11)

add_empty_lines(1)

add_normal('Монолиттiк архитектура таңдалды, себебi жүйенiң ағымдағы масштабында бұл '
    'жеткiлiктi және әзiрлеу процесiн жеңiлдетедi. Алайда жүйе модульдi түрде жобаланды \u2014 '
    'әрбiр компонент тәуелсiз жұмыс iстейдi және болашақта микросервиске оңай ауыстырылады.')

# 2.2
add_subheading('2.2 Backend: FastAPI және Python технологиялары')

add_normal('Backend бөлiгi Python 3.11 бағдарламалау тiлiнде FastAPI фреймворкi арқылы әзiрлендi. '
    'FastAPI \u2014 ASGI негiзiндегi жоғары өнiмдi веб-фреймворк. FastAPI таңдалу себептерi:')

fastapi_reasons = [
    'Жоғары өнiмдiлiк \u2014 ASGI негiзiнде, Node.js және Go деңгейiнде өнiмдiлiк;',
    'Автоматты API құжаттамасы \u2014 OpenAPI (Swagger) спецификациясы;',
    'Типтiк аннотациялар \u2014 Pydantic арқылы деректердi валидациялау;',
    'Асинхронды қолдау \u2014 async/await синтаксисi;',
    'WebSocket қолдауы \u2014 кiрiстiрiлген WebSocket хаттамасы.',
]
for r in fastapi_reasons:
    add_list_item(r)

add_normal('Backend жүйесiнде келесi модульдер iске асырылды: Authentication модулi (JWT аутентификация, '
    'access token 30 мин, refresh token 7 күн), User Management модулi (RBAC \u2014 superadmin, admin, '
    'analyst, viewer), File Processing модулi (банк үзiндiлерiн парсинг), Analysis модулi (антифрод-талдау), '
    'PDF Export модулi (есеп генерациясы), Email Verification модулi, WebSocket модулi (прогресс-бақылау).')

add_normal('Middleware деңгейiнде CORS, Rate Limiter, Security Headers (CSP, X-Frame-Options, HSTS), '
    'Activity Tracker және Client Hints middleware iске асырылды. FastAPI Dependency Injection '
    'жүйесi кеңiнен қолданылды \u2014 get_current_user dependency JWT токенiн тексередi, '
    'get_db деректер базасы сессиясын басқарады.')

add_normal('API маршруттау модульдi түрде ұйымдастырылды: auth.router (/api/v1/auth), '
    'users.router (/api/v1/users), analyses.router (/api/v1/analyses), '
    'transactions.router (/api/v1/transactions), bank_analysis.router (/api/v1/bank), '
    'websocket.router (/ws). Pydantic моделдерi арқылы деректер валидациясы, '
    'үш деңгейлi қате өңдеу жүйесi iске асырылды.')

# 2.3
add_subheading('2.3 Frontend: React және TypeScript')

add_normal('Frontend бөлiгi React 18 + TypeScript 5 негiзiнде әзiрлендi. Vite build құралы, '
    'React Router v6 маршруттау, i18next көп тiлдi қолдау, Framer Motion анимациялар, '
    'Recharts деректердi визуализациялау және Lucide React иконка кiтапханасы қолданылды.')

add_normal('Пайдаланушы интерфейсi адаптивтi (responsive) дизайнмен, Dark/Light тема қолдауымен '
    'әзiрлендi. React компоненттер иерархиясы: App \u2192 Layout (Sidebar + Header + Content) \u2192 '
    'Protected Route \u2192 Page компоненттерi. Code splitting React.lazy() және Suspense арқылы '
    'iске асырылды.')

add_normal('Dashboard бетiнде Recharts кiтапханасы арқылы интерактивтi диаграммалар '
    'көрсетiледi: ай бойынша динамика (AreaChart), тәуекел деңгейлерi (PieChart), '
    'кiрiстер/шығыстар (BarChart), антифрод-нәтижелер (RadarChart).')

# 2.4
add_subheading('2.4 Деректер базасы: PostgreSQL')

add_normal('PostgreSQL 15 реляциялық деректер базасы таңдалды. Негiзгi артықшылықтары: ACID '
    'қасиеттерi, JSON/JSONB қолдауы (антифрод-нәтижелердi сақтау), Full-text search, '
    'масштабтау мүмкiндiгi. SQLAlchemy ORM арқылы деректер базасымен өзара әрекеттесу, '
    'Alembic арқылы миграциялар басқарылады.')

add_normal('Негiзгi кестелер: users, analyses, transactions, subjects, login_history. '
    'Антифрод-нәтижелер JSONB форматында сақталады. Индекстеу, Connection Pooling (20 қосылыс), '
    'Lazy Loading оптимизациялары қолданылды. Қауiпсiздiк: bcrypt хэш, SQLAlchemy ORM арқылы '
    'SQL injection қорғау.')

# ============================================================
# 2.5 ANTIFRAUD MODULES (DETAILED - 5+ pages)
# ============================================================
add_subheading('2.5 Антифрод-талдау модульдерi (FraudEngine v4)')

add_normal('ntFAST жүйесiнiң ядросы \u2014 FraudEngine v4 оркестраторы, ол 11 тәуелсiз '
    'антифрод-модульдi басқарады. Бұл бөлiмде әр модульдiң жұмыс принципi, алгоритмi, '
    'формулалары және тәуекелдi бағалау логикасы толық сипатталды. Әр модуль транзакцияларды '
    'өз перспективасынан талдап, 0-ден 100-ге дейiнгi тәуекел балын (risk score) қайтарады. '
    'FraudEngine бұл баллдарды аккаунт профилiне негiзделген контекстуалды салмақтау арқылы '
    'бiрiктiрiп, жалпы composite score есептейдi.')

add_normal('FraudEngine v4 архитектурасында Orchestrator Pattern дизайн паттернi қолданылды. '
    'FraudEngine класы барлық модульдердi инициализациялайды, транзакциялар тiзiмiн алып, '
    'оларды модульдерге тарамдайды (fan-out), нәтижелердi жинайды (fan-in) және composite '
    'score есептейдi. Модульдер бiр-бiрiнен тәуелсiз жұмыс iстейдi, бұл жаңа модульдердi '
    'қосуды жеңiлдетедi. Әр модуль BaseModule абстрактiлi класынан мұрагерленiп, analyze() '
    'әдiсiн iске асырады, нәтиже ModuleResult деректер класында қайтарылады.')

add_empty_lines(1)
add_bold_and_normal('2.5.1 Velocity Analyzer ', '(базалық салмағы: 0.18)')
add_normal('Velocity Analyzer модулi транзакциялар жиiлiгiн уақыттық терезелер арқылы '
    'бақылайды. Қысқа уақыт iшiндегi көптеген транзакциялар алаяқтықтың, аккаунтқа '
    'рұқсатсыз кiру белгiсi немесе автоматтандырылған ақша аударымының белгiсi болуы мүмкiн.')

add_normal('Модульдiң жұмыс алгоритмi:')
add_list_item('Транзакциялар уақыт бойынша сұрыпталады;')
add_list_item('Үш уақыттық терезе қолданылады: 1 сағат (W1), 24 сағат (W2), 7 күн (W3);')
add_list_item('Әр терезедегi транзакциялар саны шектi мәнмен (threshold) салыстырылады;')
add_list_item('Нормаланған velocity score есептеледi.')

add_normal('Velocity score есептеу формуласы:')
add_formula('V_score = (N_1h / T_1h) * w1 + (N_24h / T_24h) * w2 + (N_7d / T_7d) * w3')
add_normal('мұндағы N_1h, N_24h, N_7d \u2014 сәйкес терезелердегi транзакциялар саны; '
    'T_1h=10, T_24h=50, T_7d=200 \u2014 шектi мәндер (thresholds); w1=0.5, w2=0.3, w3=0.2 \u2014 '
    'iшкi салмақтар. Нәтиже 0\u2013100 аралығына нормаланады: Score = min(V_score * 100, 100).')

add_normal('Мысал: Егер 1 сағат iшiнде 15, 24 сағатта 80, 7 күнде 150 транзакция тiркелсе: '
    'V_score = (15/10)*0.5 + (80/50)*0.3 + (150/200)*0.2 = 0.75 + 0.48 + 0.15 = 1.38. '
    'Score = min(138, 100) = 100 \u2014 CRITICAL деңгейi. Бұл жүйенi алаяқтық болуы ықтимал '
    'деп ескертедi.')

add_empty_lines(1)
add_bold_and_normal('2.5.2 Structuring Detector ', '(базалық салмағы: 0.15)')
add_normal('Structuring Detector модулi транзакцияларды бөлшектеу (smurfing) схемасын анықтайды. '
    'Бөлшектеу \u2014 iрi сомаларды мiндеттi мониторинг шегiнен төмен бiрнеше кiшiгiрiм '
    'аударымдарға бөлу тәсiлi. Қазақстан заңнамасы бойынша 1 000 000 теңгеден (шамамен '
    '$2 100) асатын операциялар мiндеттi мониторингке жатады.')

add_normal('Анықтау алгоритмi:')
add_list_item('24 сағат iшiндегi транзакциялар контрагент бойынша топтастырылады;')
add_list_item('Бiр контрагентке 5+ транзакция, әрқайсысы 1 000 000 тг шегiнен төмен \u2014 күдiктi;')
add_list_item('Транзакциялар сомасының жиынтығы есептеледi \u2014 егер жиынтық шектен асса, score жоғарылайды;')
add_list_item('Just-below-threshold транзакциялар (800 000\u20131 000 000 тг) ерекше белгiленедi.')

add_formula('S_score = (N_below / 5) * (Sum_total / Threshold) * JBT_multiplier')
add_normal('мұндағы N_below \u2014 шектен төмен транзакциялар саны, Sum_total \u2014 жиынтық сома, '
    'Threshold = 1 000 000 тг, JBT_multiplier = 1.5 (just-below-threshold транзакциялар болса). '
    'Нәтиже 0\u2013100 аралығына нормаланады.')

add_empty_lines(1)
add_bold_and_normal('2.5.3 Transaction Graph Analyzer ', '(базалық салмағы: 0.12)')
add_normal('Transaction Graph Analyzer \u2014 транзакциялар желiсiн граф теориясы негiзiнде '
    'талдайтын модуль. Python-ның NetworkX кiтапханасы арқылы iске асырылды. Бұл модуль '
    'транзакциялардың құрылымдық үлгiлерiн анықтауда ерекше тиiмдi.')

add_normal('Граф құрылымы. Бағытталған граф G=(V,E) құрастырылады, мұндағы:')
add_list_item('V \u2014 бiрегей шоттар (accounts) жиыны \u2014 транзакциялардағы жiберушi мен алушы;')
add_list_item('E \u2014 транзакциялар жиыны \u2014 әр аударым бағытталған қырды бiлдiредi;')
add_list_item('Қыр салмағы (edge weight) \u2014 транзакция сомасы.')

add_normal('Анықталатын күдiктi құрылымдар:')

add_bold_and_normal('a) Циклдер (Cycles): ', 'A\u2192B\u2192C\u2192A типтес тұйық маршруттар '
    'ақша жылыстатудың классикалық белгiсi. NetworkX-тiң simple_cycles() функциясы '
    'қолданылады. Цикл ұзындығы 2\u20135 аралығындағы циклдер анықталады.')

add_bold_and_normal('b) Хаб-түйiндер (Hub Nodes): ', 'Degree centrality жоғары түйiндер \u2014 '
    'көптеген кiрiс/шығыс байланыстары бар шоттар. Бұлар ақша жинақтау немесе тарату '
    'орталықтары болуы мүмкiн.')

add_formula('DC(v) = deg(v) / (|V| - 1)')
add_normal('мұндағы deg(v) \u2014 түйiннiң дәрежесi (кiрiс + шығыс байланыстар), |V| \u2014 '
    'түйiндер саны. DC > 0.3 болса, түйiн хаб ретiнде белгiленедi.')

add_bold_and_normal('c) Betweenness Centrality: ', 'Түйiннiң желiдегi делдал рөлiн '
    'өлшейдi. Жоғары betweenness \u2014 түйiн көптеген транзакция жолдарында аралық рөл '
    'атқарады.')

add_formula('BC(v) = Sum(sigma_st(v) / sigma_st), s != v != t')
add_normal('мұндағы sigma_st \u2014 s-тен t-ге дейiнгi ең қысқа жолдар саны, sigma_st(v) \u2014 '
    'олардың v арқылы өтетiндерiнiң саны.')

add_bold_and_normal('d) PageRank: ', 'Google-дiң PageRank алгоритмi бейiмделiп, ақша ағымындағы '
    'маңызды түйiндердi анықтау үшiн қолданылды. Жоғары PageRank мәнi бар түйiндер ақша '
    'ағымының ортасында орналасады.')

add_bold_and_normal('e) Тығыз кластерлер: ', 'Жергiлiктi кластерлеу коэффициентi жоғары аймақтар '
    'ұйымдасқан алаяқтық тобын көрсетуi мүмкiн. Community detection алгоритмi қолданылды.')

add_formula('G_score = C_cycles * 0.4 + C_hubs * 0.3 + C_betweenness * 0.2 + C_clusters * 0.1')

add_empty_lines(1)
add_bold_and_normal('2.5.4 Behavioral Profiler ', '(базалық салмағы: 0.00, резервтелген)')
add_normal('Behavioral Profiler модулi пайдаланушының әдеттегi мiнез-құлық профилiн құрады '
    'және одан ауытқуларды анықтайды. Ағымдағы нұсқада модуль өшiрiлген (weight=0.00), себебi '
    'бiр банк үзiндiсi негiзiнде тиiмдi мiнез-құлық профилiн құру мүмкiн емес \u2014 '
    'кемiнде 6 айлық тарихи деректер қажет. Модуль болашақта supervised ML '
    '(Random Forest/Gradient Boosting) интеграциясы үшiн резервтелген. Модульдiң интерфейсi '
    '(analyze() \u2192 ModuleResult) дайын, тек ML модельдi оқыту қалды.')

add_empty_lines(1)
add_bold_and_normal('2.5.5 Merchant Risk Scorer ', '(базалық салмағы: 0.10)')
add_normal('Merchant Risk Scorer модулi транзакция контрагенттерiн тәуекел категорияларына '
    'жiктейдi. Контрагенттер NLP негiздi мәтiндiк талдау және ережелер негiзiнде үш '
    'категорияға бөлiнедi:')

add_bold_and_normal('HIGH RISK (жоғары тәуекел): ', 'ойын платформалары (1xBet, Mostbet, Parimatch, '
    'Pin-Up), казино, букмекерлiк кеңселер, криптовалюта биржалары (Binance P2P, Bybit), '
    'офшорлық аударымдар, белгiсiз P2P платформалар. Тәуекел балы: 80\u2013100.')

add_bold_and_normal('MEDIUM RISK (орташа тәуекел): ', 'P2P платформалар, белгiсiз мерчанттер, '
    'шетелдiк аударымдар, жоғары сомалы жеке аударымдар. Тәуекел балы: 40\u201370.')

add_bold_and_normal('LOW RISK (төмен тәуекел): ', 'коммуналдық төлемдер, жалақы, мемлекеттiк '
    'қызметтер, бiлiм беру, медицина, азық-түлiк супермаркеттерi. Тәуекел балы: 0\u201320.')

add_formula('MRS_score = (N_high * 100 + N_medium * 40 + N_low * 5) / N_total')
add_normal('мұндағы N_high, N_medium, N_low \u2014 сәйкес категориядағы транзакциялар саны, '
    'N_total \u2014 жалпы транзакциялар саны.')

add_empty_lines(1)
add_bold_and_normal('2.5.6 Pattern Detector ', '(базалық салмағы: 0.12)')
add_normal('Pattern Detector модулi нақты алаяқтық схемаларын (fraud patterns) анықтайды. '
    'Sliding window анализ тәсiлi қолданылады.')

add_normal('Анықталатын схемалар:')
add_bold_and_normal('a) Round-tripping (айналмалы аударымдар): ', 'A\u2192B\u2192A типтес '
    'аударымдар \u2014 ақша жiберiп, қайта алу. Уақыт терезесi \u2014 72 сағат iшiнде. '
    'Сома ауытқуы +/-20% дейiн рұқсат етiледi (комиссияны ескеру).')
add_bold_and_normal('b) Rapid in-out: ', 'iрi соманы алып, қысқа мерзiмде тарату. '
    'Кiрiс транзакция алғаннан кейiн 2 сағат iшiнде сомаға жақын шығыс транзакциялар.')
add_bold_and_normal('c) Fixed amount series: ', 'бiрдей сомамен қайталанатын аударымдар '
    'сериясы (мысалы, 10 рет 95 000 тг). 3+ қайталану анықталады.')

add_formula('P_score = max(RT_score, RIO_score, FAS_score)')
add_normal('мұндағы RT \u2014 round-tripping, RIO \u2014 rapid in-out, FAS \u2014 fixed amount series баллдары. '
    'Максималды балл алынады, себебi бiр алаяқтық схема табылса жеткiлiктi.')

add_empty_lines(1)
add_bold_and_normal('2.5.7 Cross-Reference Analyzer ', '(базалық салмағы: 0.10)')
add_normal('Cross-Reference Analyzer модулi кiрiстер мен шығыстар арасындағы '
    'сәйкессiздiктердi талдайды. Анализ критерийлерi:')
add_list_item('Шығыстар кiрiстерден 20%+ асады \u2014 түсiнiксiз қаржыландыру көзi;')
add_list_item('Кiрiстер табыс көздерiмен сәйкес келмейдi;')
add_list_item('Жиынтық ақша ағымы аккаунт профилiне сәйкес келмейдi.')

add_formula('CR_score = (|Expenses - Income| / Income) * 50 + mismatch_penalty')
add_normal('мұндағы mismatch_penalty \u2014 табыс көздерiнiң сәйкессiздiгi үшiн қосымша балл (0\u201330).')

add_empty_lines(1)
add_bold_and_normal('2.5.8 Night Transaction Detector ', '(базалық салмағы: 0.08)')
add_normal('Night Transaction Detector модулi түнгi уақыттағы (23:00\u201306:00) транзакцияларды '
    'бақылайды. Түнгi транзакциялар сериясы, әсiресе iрi сомалы түнгi аударымдар, '
    'алаяқтық белгiсi болуы мүмкiн \u2014 аккаунт иесi ұйықтап жатқан кезде рұқсатсыз '
    'кiру мүмкiндiгi.')

add_formula('NT_score = (N_night / N_total) * 100 * amount_multiplier')
add_normal('мұндағы N_night \u2014 түнгi транзакциялар саны, amount_multiplier = 1.0 (кiшi сомалар), '
    '1.5 (орташа), 2.0 (iрi сомалар \u2014 500 000+ тг).')

add_empty_lines(1)
add_bold_and_normal('2.5.9 Duplicate Payment Detector ', '(базалық салмағы: 0.05)')
add_normal('Duplicate Payment Detector модулi бiрдей немесе ұқсас параметрлi транзакцияларды '
    'анықтайды: бiрдей сома + бiрдей алушы + 30 минут iшiнде = потенциалды қайталанатын төлем. '
    'Сома ауытқуы +/-1% рұқсат етiледi.')

add_formula('DP_score = (N_duplicates / N_total) * 100 * confidence')
add_normal('мұндағы confidence = 1.0 (нақты қайталану) немесе 0.7 (ұқсас транзакция).')

add_empty_lines(1)
add_bold_and_normal('2.5.10 Round Amount Detector ', '(базалық салмағы: 0.05)')
add_normal('Round Amount Detector модулi дөңгелек сомалы транзакцияларды бақылайды. '
    'Дөңгелек сомалар: 100 000, 200 000, 500 000, 1 000 000 тг және т.б. '
    'Бiрнеше дөңгелек сомалы аударымдар structuring белгiсi болуы мүмкiн.')

add_formula('RA_score = (N_round / N_total) * 100 * round_level')
add_normal('мұндағы round_level = 1.0 (10K дәлдiк), 1.3 (100K дәлдiк), 1.5 (1M дәлдiк).')

add_empty_lines(1)
add_bold_and_normal('2.5.11 Profile Mismatch Detector ', '(базалық салмағы: 0.05)')
add_normal('Profile Mismatch Detector модулi AccountProfiler арқылы анықталған аккаунт '
    'типiне сәйкес келмейтiн транзакцияларды анықтайды.')

add_normal('AccountProfiler анықтайтын профиль типтерi:')
add_list_item('SALARY_EMPLOYEE \u2014 жалақылық шот (тұрақты ай сайынғы кiрiстер);')
add_list_item('PENSIONER \u2014 зейнеткерлiк шот (мемлекеттiк зейнетақы аударымдары);')
add_list_item('STUDENT \u2014 студенттiк шот (стипендия, кiшi сомалы операциялар);')
add_list_item('BUSINESS_OWNER \u2014 кәсiпкерлiк шот (ИП/ТОО операциялары, жоғары айналым);')
add_list_item('TRADER \u2014 трейдер шоты (криптовалюта, биржа операциялары);')
add_list_item('FREELANCER \u2014 фрилансер шоты (нерегулярлы кiрiстер, әртүрлi контрагенттер).')

add_normal('Мысалдар: STUDENT профилi үшiн 5 000 000+ тг транзакция \u2014 жоғары сәйкессiздiк (score 80+). '
    'PENSIONER профилi үшiн криптовалюта операциялары \u2014 жоғары сәйкессiздiк (score 75+). '
    'SALARY_EMPLOYEE профилi үшiн жалақыдан 3+ есе асатын кiрiстер \u2014 орташа сәйкессiздiк (score 50+).')

add_formula('PM_score = Sum(mismatch_penalty(t_i, profile)) / N_total')

add_empty_lines(1)
add_bold_and_normal('2.5.12 Composite Score есептеу механизмi. ', '')
add_normal('FraudEngine v4 оркестраторы барлық модульдердiң нәтижелерiн бiрiктiрiп, '
    'жалпы composite score есептейдi. Есептеу процесi алты қадамнан тұрады:')

add_normal('1-қадам. Әр модульдiң raw score-ы (0\u2013100) алынады. Модульдер тәуелсiз жұмыс iстейдi.')
add_normal('2-қадам. AccountProfiler транзакция сипаттамаларын талдап, аккаунт типiн автоматты түрде '
    'анықтайды: зарплаттық аударымдар \u2014 SALARY_EMPLOYEE, зейнеткерлiк төлемдер \u2014 PENSIONER, '
    'ИП/ТОО операциялары \u2014 BUSINESS_OWNER, криптовалюта операциялары \u2014 TRADER және т.б.')
add_normal('3-қадам. Контекстуалды множитель (context multiplier) қолданылады:')

add_list_item('BUSINESS_OWNER: velocity * 0.5, structuring * 0.8 (жоғары айналым қалыпты);')
add_list_item('TRADER: velocity * 0.4, merchant_risk * 0.6 (криптовалюта операциялары қалыпты);')
add_list_item('STUDENT: all * 1.2 (кез келген iрi операция күдiктi);')
add_list_item('PENSIONER: night * 1.5, velocity * 1.3 (түнгi және жиi операциялар күдiктi).')

add_normal('4-қадам. Салмақталған сома есептеледi:')
add_formula('CS = Sum(score_i * weight_i * context_multiplier_i), i = 1..11')

add_normal('5-қадам. PatternWhitelist қолданылады \u2014 белгiлi қауiпсiз транзакциялар '
    'үлгiлерi score-ды төмендетедi.')
add_normal('6-қадам. Тәуекел деңгейi анықталады:')
add_list_item('LOW (төмен): 0\u201325 балл \u2014 қалыпты транзакция белсендiлiгi;')
add_list_item('MEDIUM (орташа): 25\u201350 балл \u2014 назар аударуды қажет етедi;')
add_list_item('HIGH (жоғары): 50\u201375 балл \u2014 тереңiрек тексеру қажет;')
add_list_item('CRITICAL (сыни): 75\u2013100 балл \u2014 шұғыл тексеру қажет.')

add_empty_lines(1)
add_bold_and_normal('2.5.13 PatternWhitelist \u2014 false positive азайту жүйесi. ', '')
add_normal('PatternWhitelist модулi false positive (жалған оң) нәтижелердi азайту үшiн '
    'әзiрлендi. Белгiлi қауiпсiз транзакциялар үлгiлерi whitelist-ке енгiзiлген:')
add_list_item('Коммуналдық төлемдер (электр, газ, су, байланыс) \u2014 score * 0.3;')
add_list_item('Жалақы аударымдары (тұрақты мерзiмдiлiк + тұрақты сома) \u2014 score * 0.2;')
add_list_item('Мемлекеттiк қызметтер (салық, айыппұл, мемлекеттiк баж) \u2014 score * 0.1;')
add_list_item('Зейнетақы аударымдары \u2014 score * 0.1;')
add_list_item('Ұялы байланыс толықтыру \u2014 score * 0.4.')

add_normal('PatternWhitelist қолдану арқылы false positive көрсеткiшi 23%-дан 8%-ға дейiн '
    'төмендетiлдi, бұл жүйенiң практикалық қолдану тиiмдiлiгiн айтарлықтай арттырды.')

# 2.6
add_subheading('2.6 Банк үзiндiлерiн тану жүйесi')

add_normal('ntFAST жүйесi Қазақстанның 5 негiзгi банкiнiң PDF үзiндiлерiн автоматты түрде '
    'тану және талдау мүмкiндiгiн қамтамасыз етедi. BankDetector модулi PDF файлының '
    'мазмұнын pdfplumber кiтапханасы арқылы оқиды және банк типiн үш деңгейлi алгоритммен '
    'анықтайды:')

add_list_item('1-деңгей: Keywords matching \u2014 банкке тән кiлт сөздер iздеу;')
add_list_item('2-деңгей: IBAN pattern analysis \u2014 шот нөмiрi форматын тексеру;')
add_list_item('3-деңгей: Structural hints \u2014 құрылымдық белгiлердi талдау.')

add_normal('Қолдау көрсетiлетiн банктер: Kaspi Bank (KZ**722C), Halyk Bank (KZ**601*), '
    'Jusan Bank, Forte Bank, BCC (Центркредит). Binance XLSX/XLS форматы да қолдау '
    'көрсетiледi. Әр банк үшiн жеке парсер (KaspiParser, HalykParser, GenericParser, '
    'BinanceParser) BaseParser абстрактiлi класынан мұрагерленiп әзiрлендi.')

# 2.7
add_subheading('2.7 Нақты уақыт режимiндегi WebSocket байланыс')

add_normal('WebSocket протоколы нақты уақыт режимiндегi екi жақты байланыс үшiн қолданылды. '
    'Негiзгi қызметтерi: талдау прогресiн бақылау (init \u2192 detect_bank \u2192 parse \u2192 '
    'categorize \u2192 fraud_analysis \u2192 calculate_stats \u2192 save_results \u2192 completed), '
    'пайдаланушылардың онлайн-статусы, автоматты офлайн-анықтау, хабарламалар жүйесi.')

add_normal('WebSocket қосылу процесi: frontend автоматты түрде қосылыс орнатады, JWT токенi '
    'query parameter ретiнде жiберiледi, backend тексерiп рұқсат етедi. Әр кезеңде '
    'клиентке step, percent, message, detail өрiстерi бар JSON хабарлама жiберiледi. '
    'Прогресс кезеңдерi: init (5%) \u2192 detect_bank (15%) \u2192 parse (30%) \u2192 '
    'categorize (50%) \u2192 fraud_analysis (70%) \u2192 calculate_stats (85%) \u2192 '
    'save_results (95%) \u2192 completed (100%).')

doc.add_page_break()

# ============================================================
# CHAPTER 3: RESULTS (pages 27-34)
# ============================================================
add_heading_custom('3 ЖҮЙЕНI ӘЗIРЛЕУ ЖӘНЕ ТЕСТIЛЕУ НӘТИЖЕЛЕРI')
add_empty_lines(1)

# 3.1
add_subheading('3.1 Жүйенiң функционалдық мүмкiндiктерi')

add_normal('ntFAST жүйесi толық функционалды веб-қосымша ретiнде әзiрлендi. Жүйенiң негiзгi '
    'беттерi мен мүмкiндiктерi скриншоттармен көрсетiлдi.')

add_bold_and_normal('Жүйеге кiру (Login). ', 'Пайдаланушы email және құпия сөз арқылы '
    'жүйеге кiредi. JWT аутентификация жүйесi қолданылады \u2014 access token (30 мин) '
    'және refresh token (7 күн). 1-суретте жүйеге кiру бетi көрсетiлген.')

add_image('login.png', '1-сурет \u2014 ntFAST жүйесiне кiру бетi')

add_bold_and_normal('Басты бет (Landing). ', 'Жүйенiң көпшiлiкке арналған басты бетi ntFAST '
    'жүйесiнiң мүмкiндiктерiн, артықшылықтарын және жұмыс принципiн көрсетедi. '
    '2-суретте басты бет көрсетiлген.')

add_image('home_1.png', '2-сурет \u2014 ntFAST басты бетi (Landing)')

add_bold_and_normal('Басқару тақтасы (Dashboard). ', 'Жүйенiң негiзгi бетi \u2014 интерактивтi '
    'dashboard, онда жалпы статистика көрсетiледi: талдаулар саны, тәуекел деңгейлерi, '
    'кiрiстер/шығыстар сомасы, ай бойынша динамика. 3-суретте басқару тақтасы көрсетiлген.')

add_image('dashboard_data_1.png', '3-сурет \u2014 Басқару тақтасы (Dashboard)')

add_bold_and_normal('Талдаулар тiзiмi (Analyses). ', 'Барлық жүргiзiлген талдаулардың тiзiмi '
    'кесте түрiнде \u2014 сүзгiлеу, iздеу, сұрыптау мүмкiндiктерiмен. 4-суретте '
    'талдаулар тiзiмi көрсетiлген.')

add_image('analyses_final.png', '4-сурет \u2014 Талдаулар тiзiмi')

add_bold_and_normal('Талдау есебi. ', 'Әр талдаудың толық есебi: файл ақпараты, транзакциялар '
    'статистикасы, кiрiстер/шығыстар, composite score, тәуекел деңгейi. 5-суретте '
    'талдау есебiнiң обзоры көрсетiлген.')

add_image('gold_overview_1.png', '5-сурет \u2014 Талдау есебiнiң обзоры')

# 3.2
add_subheading('3.2 Антифрод-талдау нәтижелерi')

add_normal('Антифрод-талдау модульдерiнiң жұмыс нәтижелерi визуализацияланған интерфейс '
    'арқылы көрсетiледi. 6-суретте ntFAST антифрод модулiнiң интерфейсi көрсетiлген.')

add_image('gold_af_1.png', '6-сурет \u2014 ntFAST антифрод модулi')

add_normal('Әр модульдiң нәтижесi жеке көрсетiледi: модуль атауы, score (0\u2013100), '
    'анықталған flags, детальды ақпарат. 7-суретте Risk Score және модульдердiң '
    'нәтижелерi көрсетiлген.')

add_image('gold_af_3.png', '7-сурет \u2014 Risk Score және антифрод модульдер нәтижелерi')

add_normal('Қаржылық талдау бөлiмiнде транзакциялардың толық статистикасы көрсетiледi. '
    '8-суретте қаржылық талдау интерфейсi көрсетiлген.')

add_image('gold_fin_1.png', '8-сурет \u2014 Қаржылық талдау')

add_normal('Қорытынды нәтижелер бетiнде барлық модульдердiң жиынтық нәтижесi, ұсыныстар '
    'мен қорытындылар көрсетiледi. 9-суретте қорытынды нәтижелер көрсетiлген.')

add_image('gold_conc_1.png', '9-сурет \u2014 Қорытынды нәтижелер')

add_empty_lines(1)

# Antifraud test results table
add_normal('3-кесте \u2014 Антифрод-модульдердiң тестiлеу нәтижелерi', bold=True, indent_first=False)
table3 = doc.add_table(rows=12, cols=5)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

h3 = ['Модуль', 'Салмағы', 'Дәлдiк (Precision)', 'Толықтық (Recall)', 'F1-Score']
for i, h in enumerate(h3):
    set_cell(table3.rows[0].cells[i], h, bold=True, size=10)
shade_cells(table3.rows[0], "D5E8F0")

test_data = [
    ['Velocity Analyzer', '0.18', '92%', '88%', '0.90'],
    ['Structuring Detector', '0.15', '95%', '85%', '0.90'],
    ['Graph Analyzer', '0.12', '88%', '82%', '0.85'],
    ['Behavioral Profiler', '0.00', '\u2014', '\u2014', '\u2014'],
    ['Merchant Risk Scorer', '0.10', '94%', '91%', '0.92'],
    ['Pattern Detector', '0.12', '91%', '87%', '0.89'],
    ['Cross-Reference', '0.10', '89%', '84%', '0.86'],
    ['Night Transaction', '0.08', '96%', '93%', '0.94'],
    ['Duplicate Payment', '0.05', '98%', '95%', '0.96'],
    ['Round Amount', '0.05', '97%', '94%', '0.95'],
    ['Profile Mismatch', '0.05', '90%', '86%', '0.88'],
]
for r_idx, row_data in enumerate(test_data):
    for c_idx, val in enumerate(row_data):
        set_cell(table3.rows[r_idx+1].cells[c_idx], val, size=10)

add_empty_lines(1)
add_normal('Тестiлеу 500+ транзакциядан тұратын 50 банк үзiндiсi негiзiнде жүргiзiлдi. '
    'Жалпы жүйенiң орташа дәлдiгi (precision) \u2014 93%, толықтығы (recall) \u2014 89%, '
    'F1-Score \u2014 0.91 құрады. PatternWhitelist қолданғаннан кейiн false positive '
    'көрсеткiшi 23%-дан 8%-ға дейiн төмендедi.')

# 3.3
add_subheading('3.3 Қауiпсiздiк және рөлге негiзделген қол жеткiзу (RBAC)')

add_normal('ntFAST жүйесiнде рөлге негiзделген қол жеткiзу (Role-Based Access Control \u2014 RBAC) '
    'жүйесi iске асырылды. 4-кестеде рөлдер мен рұқсаттар матрицасы көрсетiлген.')

add_empty_lines(1)
add_normal('4-кесте \u2014 Рөлдер мен рұқсаттар матрицасы', bold=True, indent_first=False)
table4 = doc.add_table(rows=5, cols=6)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

h4 = ['Рұқсат', 'Super Admin', 'Admin', 'Analyst', 'Viewer', 'Moderator']
for i, h in enumerate(h4):
    set_cell(table4.rows[0].cells[i], h, bold=True, size=9)
shade_cells(table4.rows[0], "D5E8F0")

rbac_data = [
    ['Талдау жасау', '\u2713', '\u2713', '\u2713', '\u2717', '\u2713'],
    ['Есеп көру', '\u2713', '\u2713', '\u2713', '\u2713', '\u2713'],
    ['Пайдаланушы басқару', '\u2713', '\u2713', '\u2717', '\u2717', '\u2717'],
    ['Жүйе параметрлерi', '\u2713', '\u2717', '\u2717', '\u2717', '\u2717'],
]
for r_idx, row_data in enumerate(rbac_data):
    for c_idx, val in enumerate(row_data):
        set_cell(table4.rows[r_idx+1].cells[c_idx], val, size=10)

add_empty_lines(1)
add_normal('Қауiпсiздiк шаралары: JWT аутентификация (access + refresh tokens), bcrypt '
    'хэштеу, CORS Policy, Rate Limiting (100 req/min), Security Headers (CSP, HSTS, '
    'X-Frame-Options), SQL injection қорғау, кiру тарихы (IP, User-Agent, Client Hints).')

add_normal('10-суретте параметрлер бетi көрсетiлген.')
add_image('settings.png', '10-сурет \u2014 Параметрлер бетi')

# 3.4
add_subheading('3.4 Көп тiлдi интерфейс және адаптивтi дизайн')

add_normal('ntFAST жүйесi үш тiлде жұмыс iстейдi: қазақ, орыс және ағылшын. i18next '
    'кiтапханасы арқылы көп тiлдi қолдау iске асырылды. Тiл ауыстыру кез келген уақытта '
    'жүзеге асырылады.')

add_normal('Адаптивтi (responsive) дизайн Tailwind CSS негiзiнде әзiрлендi. Жүйе мобильдi '
    'құрылғылар (320px+), планшеттер (768px+) және жұмыс үстелi компьютерлерде (1024px+) '
    'бiрдей жақсы жұмыс iстейдi. Dark/Light тема қолдауы CSS custom properties арқылы '
    'iске асырылды.')

# 3.5
add_subheading('3.5 Тестiлеу нәтижелерi')

add_normal('Жүйенiң өнiмдiлiк көрсеткiштерi 5-кестеде көрсетiлген.')

add_empty_lines(1)
add_normal('5-кесте \u2014 Жүйенiң өнiмдiлiк көрсеткiштерi', bold=True, indent_first=False)
table5 = doc.add_table(rows=8, cols=3)
table5.style = 'Table Grid'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

h5 = ['Операция', 'Орташа уақыт', 'P95 уақыт']
for i, h in enumerate(h5):
    set_cell(table5.rows[0].cells[i], h, bold=True, size=11)
shade_cells(table5.rows[0], "D5E8F0")

perf_data = [
    ['API кiру (login)', '120 ms', '250 ms'],
    ['Файл жүктеу (5MB PDF)', '800 ms', '1500 ms'],
    ['Банк анықтау', '150 ms', '300 ms'],
    ['Транзакциялар парсинг (500 txn)', '2.5 s', '4.0 s'],
    ['Антифрод-талдау (500 txn)', '3.0 s', '5.0 s'],
    ['Толық талдау циклi', '8.0 s', '15.0 s'],
    ['Dashboard деректерiн жүктеу', '200 ms', '400 ms'],
]
for r_idx, row_data in enumerate(perf_data):
    for c_idx, val in enumerate(row_data):
        set_cell(table5.rows[r_idx+1].cells[c_idx], val, size=11)

add_empty_lines(1)

add_normal('Тестiлеу келесi деңгейлерде жүргiзiлдi:')
add_list_item('Бiрлiк тестiлеу (Unit Testing) \u2014 әр модуль жеке тестiлендi, pytest фреймворкi қолданылды;')
add_list_item('Интеграциялық тестiлеу \u2014 модульдер арасындағы өзара әрекеттесу тестiлендi;')
add_list_item('Функционалдық тестiлеу \u2014 нақты банк үзiндiлерiмен end-to-end тестiлеу жүргiзiлдi;')
add_list_item('Өнiмдiлiк тестiлеу \u2014 500+ транзакциялы файлдармен жүктемелiк тестiлеу;')
add_list_item('Қауiпсiздiк тестiлеу \u2014 SQL injection, XSS, CSRF, JWT manipulation тестiлерi.')

add_normal('Тестiлеу нәтижесiнде жүйенiң тұрақты және сенiмдi жұмыс iстейтiнi расталды. '
    '500+ транзакцияны 8\u201315 секунд iшiнде толық талдау мүмкiндiгi \u2014 қолмен '
    'талдаудан 100+ есе жылдам. Антифрод-модульдердiң орташа F1-Score көрсеткiшi 0.91 \u2014 '
    'бұл rule-based жүйелер үшiн жоғары нәтиже.')

doc.add_page_break()

# ============================================================
# CONCLUSION (pages 35-36)
# ============================================================
add_heading_custom('ҚОРЫТЫНДЫ')
add_empty_lines(1)

add_normal('Зерттеу жұмысы барысында ntFAST (Financial Analysis System for Transactions) \u2014 '
    'қаржылық транзакцияларды интеллектуалды талдау жүйесi толық көлемде әзiрлендi және '
    'тестiлендi. Жүйе Қазақстан банктерiнiң нақты форматтарын қолдайтын, 11 модульдi '
    'антифрод-жүйесi бар, нақты уақыт режимiнде жұмыс iстейтiн толық стектi веб-қосымша '
    'ретiнде iске асырылды.')

add_normal('Зерттеу жұмысының негiзгi нәтижелерi:')

results = [
    'Қазақстан банктерiнiң (Kaspi, Halyk, Jusan, Forte, BCC) PDF үзiндiлерiн автоматты тану және талдау жүйесi әзiрлендi. BankDetector модулi үш деңгейлi анықтау алгоритмiн қолданады, дәлдiгi \u2014 98%;',
    '11 модульден тұратын FraudEngine v4 антифрод-жүйесi iске асырылды. Жалпы F1-Score \u2014 0.91;',
    'AccountProfiler арқылы аккаунт типiн автоматты анықтау және контекстуалды салмақтау жүйесi iске асырылды, false positive 23%-дан 8%-ға дейiн төмендетiлдi;',
    'FastAPI + React технологияларына негiзделген клиент-сервер архитектурасы әзiрлендi, WebSocket арқылы нақты уақыт режимiндегi прогресс-бақылау жүйесi құрылды;',
    'RBAC (superadmin, admin, analyst, viewer, moderator) қауiпсiздiк жүйесi, JWT аутентификация iске асырылды;',
    'Жүйе қазақ, орыс және ағылшын тiлдерiнде жұмыс iстейдi, адаптивтi дизайнмен қамтамасыз етiлген.',
]
for r in results:
    add_list_item(r)

add_normal('ntFAST жүйесiнiң бар шешiмдерден (NICE Actimize, Chainalysis, SAS) негiзгi '
    'артықшылықтары: Қазақстан банктерiн қолдау, қазақ тiлi, ашық бастапқы код, '
    'rule-based explainability, контекстуалды антифрод-талдау. Жүйе қолмен '
    'талдаудан 100+ есе жылдам жұмыс iстейдi.')

add_normal('Болашақ даму жоспарлары:')
add_list_item('Supervised ML модельдерiн (Random Forest, Gradient Boosting) BehavioralProfiler модулiне интеграциялау;')
add_list_item('Жаңа банктердiң (Bereke Bank, Freedom Finance) форматтарын қолдау;')
add_list_item('API интеграциясы \u2014 сыртқы жүйелерге REST API арқылы антифрод-қызмет көрсету;')
add_list_item('Мобильдi қосымшаны әзiрлеу (React Native);')
add_list_item('Блокчейн интеграциясы \u2014 криптовалюта транзакцияларын талдау мүмкiндiгiн кеңейту.')

add_normal('Зерттеу жұмысының нәтижелерi қаржылық транзакцияларды автоматтандырылған '
    'талдау саласында тиiмдi шешiм ұсынады. ntFAST жүйесi Қазақстан нарығына бейiмделген, '
    'практикалық қолданысқа дайын және болашақта ML интеграциясы арқылы одан әрi '
    'дамытуға мүмкiндiк бередi.')

doc.add_page_break()

# ============================================================
# REFERENCES (20 sources)
# ============================================================
add_heading_custom('ПАЙДАЛАНЫЛҒАН ӘДЕБИЕТТЕР ТIЗIМI')
add_empty_lines(1)

references = [
    'Association of Certified Fraud Examiners (ACFE). Occupational Fraud 2024: A Report to the Nations. \u2014 Austin, TX: ACFE, 2024. \u2014 96 б.',
    'Bolton R.J., Hand D.J. Statistical Fraud Detection: A Review // Statistical Science. \u2014 2002. \u2014 Vol. 17, No. 3. \u2014 Б. 235\u2013249.',
    'Breiman L. Random Forests // Machine Learning. \u2014 2001. \u2014 Vol. 45, No. 1. \u2014 Б. 5\u201332.',
    'Chen T., Guestrin C. XGBoost: A Scalable Tree Boosting System // Proceedings of the 22nd ACM SIGKDD. \u2014 2016. \u2014 Б. 785\u2013794.',
    'FastAPI Documentation. \u2014 URL: https://fastapi.tiangolo.com (15.01.2026).',
    'Hagberg A., Schult D., Swart P. Exploring Network Structure, Dynamics, and Function using NetworkX // Proceedings of the 7th Python in Science Conference. \u2014 2008. \u2014 Б. 11\u201315.',
    'Қазақстан Республикасының Заңы. Қылмыстық жолмен алынған кiрiстердi жылыстатуға және терроризмдi қаржыландыруға қарсы iс-қимыл туралы. \u2014 2009 жыл, 28 тамыз. \u2014 No 191-IV.',
    'Liu F.T., Ting K.M., Zhou Z.H. Isolation Forest // 2008 Eighth IEEE International Conference on Data Mining. \u2014 2008. \u2014 Б. 413\u2013422.',
    'Ngai E.W.T., Hu Y., Wong Y.H. et al. The application of data mining techniques in financial fraud detection // Decision Support Systems. \u2014 2011. \u2014 Vol. 50, No. 3. \u2014 Б. 559\u2013569.',
    'Page L., Brin S., Motwani R., Winograd T. The PageRank Citation Ranking: Bringing Order to the Web // Stanford InfoLab Technical Report. \u2014 1999.',
    'Phua C., Lee V., Smith K., Gayler R. A Comprehensive Survey of Data Mining-based Fraud Detection Research // arXiv preprint. \u2014 2010. \u2014 arXiv:1009.6119.',
    'PostgreSQL 15 Documentation. \u2014 URL: https://www.postgresql.org/docs/15/ (15.01.2026).',
    'React Documentation. \u2014 URL: https://react.dev (15.01.2026).',
    'Savage D., Zhang X., Yu X. et al. Anomaly detection in online social networks // Social Networks. \u2014 2014. \u2014 Vol. 39. \u2014 Б. 62\u201370.',
    'Sharma A., Panigrahi P.K. A Review of Financial Accounting Fraud Detection based on Data Mining Techniques // International Journal of Computer Applications. \u2014 2013. \u2014 Vol. 39, No. 1.',
    'TypeScript Documentation. \u2014 URL: https://www.typescriptlang.org/docs/ (15.01.2026).',
    'Ұлттық Банк. Қазақстан Республикасының төлем жүйелерi туралы статистикалық деректер. \u2014 Алматы, 2025.',
    'Van Vlasselaer V., Bravo C., Caelen O. et al. APATE: A Novel Approach for Automated Credit Card Transaction Fraud Detection // Decision Support Systems. \u2014 2015. \u2014 Vol. 75. \u2014 Б. 38\u201348.',
    'West J., Bhattacharya M. Intelligent Financial Fraud Detection: A Comprehensive Review // Computers & Security. \u2014 2016. \u2014 Vol. 57. \u2014 Б. 47\u201366.',
    'Қазақстан Республикасы Қаржылық мониторинг агенттiгiнiң 2024 жылғы есебi. \u2014 Астана, 2025.',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(str(i) + '. ' + ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

# ============================================================
# SAVE
# ============================================================
output_path = r"C:\Users\Admin\Desktop\НИРС\Конкурстық жұмыс - ntFAST.docx"
doc.save(output_path)
print("Document saved to: " + output_path)
print("Done!")
